"""
File Conversion Service
Handles PDF to JPG and JPG to PDF conversions
"""

import os
import asyncio
from PIL import Image
import PyPDF2
from pdf2image import convert_from_path
import tempfile

async def convert_pdf_to_jpg(pdf_path: str) -> str:
    """
    Convert PDF to JPG
    Returns the path to the converted JPG file
    """
    try:
        # Convert first page of PDF to image
        images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=150)

        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.jpg')
        os.close(output_fd)

        # Save the image
        images[0].save(output_path, 'JPEG', quality=95)

        return output_path
    except Exception as e:
        raise ValueError(f"PDF to JPG conversion failed: {str(e)}")

async def convert_jpg_to_pdf(jpg_path: str) -> str:
    """
    Convert JPG to PDF
    Returns the path to the converted PDF file
    """
    try:
        # Open the image
        image = Image.open(jpg_path)

        # Convert RGBA to RGB if necessary
        if image.mode in ('RGBA', 'LA', 'P'):
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = rgb_image

        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.pdf')
        os.close(output_fd)

        # Save as PDF
        image.save(output_path, 'PDF')

        return output_path
    except Exception as e:
        raise ValueError(f"JPG to PDF conversion failed: {str(e)}")

# Note: Install required packages:
# pip install PyPDF2 Pillow pdf2image
# For pdf2image, also need: apt-get install poppler-utils (on Linux)

# Additional imports for new conversions
import shutil
import platform

# Try to import pdf2docx and docx2pdf
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    import docx2pdf
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


# ==================== PDF ↔ DOCX Conversions ====================

async def convert_pdf_to_docx(pdf_path: str) -> str:
    """
    Convert PDF to DOCX using pdf2docx library
    Returns the path to the converted DOCX file
    """
    if not PDF2DOCX_AVAILABLE:
        raise ValueError("pdf2docx library not available. Please install: pip install pdf2docx")
    
    try:
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.docx')
        os.close(output_fd)
        
        # Create converter and convert
        cv = Converter(pdf_path)
        cv.convert(output_path)
        cv.close()
        
        return output_path
    except Exception as e:
        raise ValueError(f"PDF to DOCX conversion failed: {str(e)}")


async def convert_docx_to_pdf(docx_path: str) -> str:
    """
    Convert DOCX to PDF
    Tries docx2pdf first (requires MS Word on Windows), falls back to python-docx + reportlab
    Returns the path to the converted PDF file
    """
    if not DOCX2PDF_AVAILABLE:
        raise ValueError("Required libraries not available. Please install: pip install docx2pdf python-docx reportlab")
    
    try:
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.pdf')
        os.close(output_fd)
        
        # Try docx2pdf first (best quality, requires Word on Windows)
        if platform.system() == 'Windows':
            try:
                docx2pdf.convert(docx_path, output_path)
                return output_path
            except Exception as e:
                print(f"docx2pdf failed, falling back to python-docx: {e}")
        
        # Fallback: Use python-docx + reportlab (cross-platform, basic formatting only)
        doc = Document(docx_path)
        pdf = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Extract text from DOCX paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                p = Paragraph(para.text, styles['Normal'])
                story.append(p)
        
        # Build PDF
        if story:
            pdf.build(story)
        else:
            raise ValueError("No content found in DOCX file")
        
        return output_path
    except Exception as e:
        raise ValueError(f"DOCX to PDF conversion failed: {str(e)}")


async def convert_doc_to_docx(doc_path: str) -> str:
    """
    Convert Word (DOC) to DOCX using LibreOffice
    Returns the path to the converted DOCX file
    """
    return await convert_with_libreoffice(doc_path, 'docx')


# ==================== Image Format Conversions ====================

async def convert_jpg_to_png(jpg_path: str) -> str:
    """
    Convert JPG to PNG
    Returns the path to the converted PNG file
    """
    try:
        # Open the image
        image = Image.open(jpg_path)
        
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.png')
        os.close(output_fd)
        
        # Save as PNG
        image.save(output_path, 'PNG')
        
        return output_path
    except Exception as e:
        raise ValueError(f"JPG to PNG conversion failed: {str(e)}")


async def convert_png_to_jpg(png_path: str) -> str:
    """
    Convert PNG to JPG
    Returns the path to the converted JPG file
    """
    try:
        # Open the image
        image = Image.open(png_path)
        
        # Convert RGBA to RGB if necessary (PNG supports transparency, JPG doesn't)
        if image.mode in ('RGBA', 'LA', 'P'):
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = rgb_image
        
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.jpg')
        os.close(output_fd)
        
        # Save as JPG
        image.save(output_path, 'JPEG', quality=95)
        
        return output_path
    except Exception as e:
        raise ValueError(f"PNG to JPG conversion failed: {str(e)}")


async def convert_jpeg_to_png(jpeg_path: str) -> str:
    """
    Convert JPEG to PNG
    Returns the path to the converted PNG file
    """
    try:
        # Open the image
        image = Image.open(jpeg_path)
        
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.png')
        os.close(output_fd)
        
        # Save as PNG
        image.save(output_path, 'PNG')
        
        return output_path
    except Exception as e:
        raise ValueError(f"JPEG to PNG conversion failed: {str(e)}")


async def convert_png_to_jpeg(png_path: str) -> str:
    """
    Convert PNG to JPEG
    Returns the path to the converted JPEG file
    """
    try:
        # Open the image
        image = Image.open(png_path)
        
        # Convert RGBA to RGB if necessary
        if image.mode in ('RGBA', 'LA', 'P'):
            rgb_image = Image.new('RGB', image.size, (255, 255, 255))
            rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
            image = rgb_image
        
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.jpeg')
        os.close(output_fd)
        
        # Save as JPEG
        image.save(output_path, 'JPEG', quality=95)
        
        return output_path
    except Exception as e:
        raise ValueError(f"PNG to JPEG conversion failed: {str(e)}")


# ==================== JPG ↔ JPEG (File Rename) ====================

async def convert_jpg_to_jpeg(jpg_path: str) -> str:
    """
    Convert JPG to JPEG (same format, just rename file extension)
    Returns the path to the renamed file
    """
    try:
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.jpeg')
        os.close(output_fd)
        
        # Simply copy the file with new extension
        shutil.copy2(jpg_path, output_path)
        
        return output_path
    except Exception as e:
        raise ValueError(f"JPG to JPEG conversion failed: {str(e)}")


async def convert_jpeg_to_jpg(jpeg_path: str) -> str:
    """
    Convert JPEG to JPG (same format, just rename file extension)
    Returns the path to the renamed file
    """
    try:
        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.jpg')
        os.close(output_fd)

        # Simply copy the file with new extension
        shutil.copy2(jpeg_path, output_path)

        return output_path
    except Exception as e:
        raise ValueError(f"JPEG to JPG conversion failed: {str(e)}")


# ==================== PDF/DOCX to Image Conversions (Multi-page support) ====================

import zipfile
import io

async def convert_pdf_to_images(pdf_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert all pages of a PDF to images
    Returns tuple: (list of image paths, is_multipage)
    For multi-page PDFs, returns multiple images
    """
    try:
        # Convert all pages of PDF to images
        images = convert_from_path(pdf_path, dpi=dpi)

        output_paths = []
        suffix = f'.{format.lower()}'
        save_format = 'PNG' if format.lower() == 'png' else 'JPEG'

        for i, image in enumerate(images):
            output_fd, output_path = tempfile.mkstemp(suffix=suffix)
            os.close(output_fd)

            # Convert RGBA to RGB for JPEG
            if save_format == 'JPEG' and image.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', image.size, (255, 255, 255))
                rgb_image.paste(image, mask=image.split()[-1] if image.mode == 'RGBA' else None)
                image = rgb_image

            image.save(output_path, save_format, quality=95 if save_format == 'JPEG' else None)
            output_paths.append(output_path)

        return output_paths, len(images) > 1
    except Exception as e:
        raise ValueError(f"PDF to {format.upper()} conversion failed: {str(e)}")


async def convert_docx_to_images(docx_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert all pages of a DOCX to images
    First converts DOCX to PDF, then PDF to images
    Returns tuple: (list of image paths, is_multipage)
    """
    try:
        # First convert DOCX to PDF
        pdf_path = await convert_docx_to_pdf(docx_path)

        try:
            # Then convert PDF to images
            return await convert_pdf_to_images(pdf_path, format, dpi)
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
    except Exception as e:
        raise ValueError(f"DOCX to {format.upper()} conversion failed: {str(e)}")


def create_images_zip(image_paths: list, format: str, base_name: str = 'converted') -> bytes:
    """
    Create a ZIP file containing multiple images
    Returns the ZIP file as bytes
    """
    zip_buffer = io.BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for i, path in enumerate(image_paths):
            with open(path, 'rb') as f:
                zip_file.writestr(f'{base_name}_page_{i + 1}.{format}', f.read())

    zip_buffer.seek(0)
    return zip_buffer.getvalue()


# ==================== Image to DOCX Conversion ====================

async def convert_image_to_docx(image_path: str) -> str:
    """
    Convert an image to DOCX by embedding the image in the document
    Returns the path to the converted DOCX file
    """
    if not DOCX2PDF_AVAILABLE:
        raise ValueError("Required libraries not available. Please install: pip install python-docx")

    try:
        from docx.shared import Inches

        # Open the image to get dimensions
        image = Image.open(image_path)
        width, height = image.size

        # Create a new DOCX document
        doc = Document()

        # Calculate appropriate width (max 6.5 inches for letter size with margins)
        max_width = 6.5
        doc_width = min(max_width, width / 96)  # Assuming 96 DPI

        # Add the image to the document, scaled appropriately
        doc.add_picture(image_path, width=Inches(doc_width))

        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.docx')
        os.close(output_fd)

        # Save the document
        doc.save(output_path)

        return output_path
    except Exception as e:
        raise ValueError(f"Image to DOCX conversion failed: {str(e)}")


# ==================== Batch Image Conversions (Multiple images to single document) ====================

async def convert_images_to_single_pdf(image_paths: list[str]) -> str:
    """
    Combine multiple images into a single PDF with one page per image.
    Handles mixed image sizes and orientations automatically.

    Args:
        image_paths: List of paths to image files (jpg, jpeg, png)

    Returns:
        Path to the combined PDF file
    """
    if not image_paths or len(image_paths) < 2:
        raise ValueError("At least 2 images are required to create a batch PDF")

    try:
        # Open all images
        images = []
        for image_path in image_paths:
            try:
                image = Image.open(image_path)

                # Convert RGBA to RGB if necessary
                if image.mode in ('RGBA', 'LA', 'P'):
                    rgb_image = Image.new('RGB', image.size, (255, 255, 255))
                    # Handle different image modes for pasting
                    if image.mode == 'RGBA':
                        rgb_image.paste(image, mask=image.split()[-1])
                    elif image.mode == 'LA':
                        rgb_image.paste(image, mask=image.split()[-1])
                    else:  # Palette mode
                        rgb_image.paste(image)
                    image = rgb_image
                elif image.mode != 'RGB':
                    # Convert grayscale or other modes to RGB
                    image = image.convert('RGB')

                images.append(image)
            except Exception as e:
                raise ValueError(f"Failed to open image {image_path}: {str(e)}")

        if not images:
            raise ValueError("No valid images found to convert")

        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.pdf')
        os.close(output_fd)

        # Save all images as a single multi-page PDF
        # Pillow's save_all parameter creates a multi-page PDF
        images[0].save(
            output_path,
            'PDF',
            save_all=True,
            append_images=images[1:] if len(images) > 1 else []
        )

        return output_path
    except Exception as e:
        raise ValueError(f"Batch images to PDF conversion failed: {str(e)}")


async def convert_images_to_single_docx(image_paths: list[str]) -> str:
    """
    Combine multiple images into a single DOCX with one page per image.
    Each image is placed on its own page.

    Args:
        image_paths: List of paths to image files (jpg, jpeg, png)

    Returns:
        Path to the combined DOCX file
    """
    if not DOCX2PDF_AVAILABLE:
        raise ValueError("Required libraries not available. Please install: pip install python-docx")

    if not image_paths or len(image_paths) < 2:
        raise ValueError("At least 2 images are required to create a batch DOCX")

    try:
        from docx.shared import Inches

        # Create a new DOCX document
        doc = Document()

        # Add each image on its own page
        for i, image_path in enumerate(image_paths):
            try:
                # Open the image to get dimensions
                image = Image.open(image_path)
                width, height = image.size

                # Calculate appropriate width (max 6.5 inches for letter size with margins)
                max_width = 6.5
                doc_width = min(max_width, width / 96)  # Assuming 96 DPI

                # Add the image to the document, scaled appropriately
                doc.add_picture(image_path, width=Inches(doc_width))

                # Add page break after each image except the last one
                if i < len(image_paths) - 1:
                    doc.add_page_break()
            except Exception as e:
                raise ValueError(f"Failed to process image {image_path}: {str(e)}")

        # Create temporary file for output
        output_fd, output_path = tempfile.mkstemp(suffix='.docx')
        os.close(output_fd)

        # Save the document
        doc.save(output_path)

        return output_path
    except Exception as e:
        raise ValueError(f"Batch images to DOCX conversion failed: {str(e)}")


# ==================== LibreOffice-based Conversions (Excel, PowerPoint) ====================

import subprocess

def _check_libreoffice_available() -> bool:
    """Check if LibreOffice is installed and available"""
    try:
        result = subprocess.run(
            ['libreoffice', '--version'],
            capture_output=True,
            text=True,
            timeout=10
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


async def convert_with_libreoffice(input_path: str, output_format: str) -> str:
    """
    Convert a file using LibreOffice headless mode
    Supports: xlsx/xls -> pdf, pptx/ppt -> pdf, docx -> pdf, etc.
    Returns the path to the converted file
    """
    if not _check_libreoffice_available():
        raise ValueError("LibreOffice is not installed. Please install: sudo apt install libreoffice")

    try:
        # Create temp directory for output
        output_dir = tempfile.mkdtemp()

        # Run LibreOffice conversion
        loop = asyncio.get_event_loop()
        result = await loop.run_in_executor(
            None,
            lambda: subprocess.run(
                [
                    'libreoffice',
                    '--headless',
                    '--convert-to', output_format,
                    '--outdir', output_dir,
                    input_path
                ],
                capture_output=True,
                text=True,
                timeout=120  # 2 minute timeout for large files
            )
        )

        if result.returncode != 0:
            raise ValueError(f"LibreOffice conversion failed: {result.stderr}")

        # Find the output file (LibreOffice names it based on input filename)
        input_basename = os.path.splitext(os.path.basename(input_path))[0]
        output_filename = f"{input_basename}.{output_format}"
        output_path = os.path.join(output_dir, output_filename)

        if not os.path.exists(output_path):
            # Sometimes LibreOffice outputs with different extension
            for f in os.listdir(output_dir):
                if f.endswith(f'.{output_format}'):
                    output_path = os.path.join(output_dir, f)
                    break

        if not os.path.exists(output_path):
            raise ValueError(f"Conversion completed but output file not found")

        # Move to a proper temp file
        final_fd, final_path = tempfile.mkstemp(suffix=f'.{output_format}')
        os.close(final_fd)
        shutil.move(output_path, final_path)

        # Clean up temp directory
        shutil.rmtree(output_dir, ignore_errors=True)

        return final_path
    except subprocess.TimeoutExpired:
        raise ValueError("Conversion timed out. The file may be too large or complex.")
    except Exception as e:
        raise ValueError(f"LibreOffice conversion failed: {str(e)}")


# ==================== Excel Conversions ====================

async def convert_xlsx_to_pdf(xlsx_path: str) -> str:
    """
    Convert Excel (XLSX) to PDF using LibreOffice
    Returns the path to the converted PDF file
    """
    return await convert_with_libreoffice(xlsx_path, 'pdf')


async def convert_xls_to_pdf(xls_path: str) -> str:
    """
    Convert Excel (XLS) to PDF using LibreOffice
    Returns the path to the converted PDF file
    """
    return await convert_with_libreoffice(xls_path, 'pdf')


async def convert_xls_to_xlsx(xls_path: str) -> str:
    """
    Convert Excel (XLS) to XLSX using LibreOffice
    Returns the path to the converted XLSX file
    """
    return await convert_with_libreoffice(xls_path, 'xlsx')


async def convert_xlsx_to_docx(xlsx_path: str) -> str:
    """
    Convert Excel (XLSX) to DOCX using LibreOffice
    Returns the path to the converted DOCX file
    """
    return await convert_with_libreoffice(xlsx_path, 'docx')


async def convert_pdf_to_xlsx(pdf_path: str) -> str:
    """
    Convert PDF to Excel (XLSX) using LibreOffice
    Note: Best results with PDFs that contain tabular data
    Returns the path to the converted XLSX file
    """
    return await convert_with_libreoffice(pdf_path, 'xlsx')


async def convert_xls_to_images(xls_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert Excel (XLS) to images
    First converts to PDF, then PDF to images
    Returns tuple: (list of image paths, is_multipage)
    """
    try:
        # First convert XLS to PDF
        pdf_path = await convert_xls_to_pdf(xls_path)

        try:
            # Then convert PDF to images
            return await convert_pdf_to_images(pdf_path, format, dpi)
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
    except Exception as e:
        raise ValueError(f"XLS to {format.upper()} conversion failed: {str(e)}")


# ==================== PowerPoint Conversions ====================

async def convert_pptx_to_pdf(pptx_path: str) -> str:
    """
    Convert PowerPoint (PPTX) to PDF using LibreOffice
    Returns the path to the converted PDF file
    """
    return await convert_with_libreoffice(pptx_path, 'pdf')


async def convert_ppt_to_pdf(ppt_path: str) -> str:
    """
    Convert PowerPoint (PPT) to PDF using LibreOffice
    Returns the path to the converted PDF file
    """
    return await convert_with_libreoffice(ppt_path, 'pdf')


async def convert_ppt_to_pptx(ppt_path: str) -> str:
    """
    Convert PowerPoint (PPT) to PPTX using LibreOffice
    Returns the path to the converted PPTX file
    """
    return await convert_with_libreoffice(ppt_path, 'pptx')


async def convert_pdf_to_pptx(pdf_path: str) -> str:
    """
    Convert PDF to PowerPoint (PPTX) using LibreOffice
    Note: Results may vary depending on PDF content
    Returns the path to the converted PPTX file
    """
    return await convert_with_libreoffice(pdf_path, 'pptx')


async def convert_pptx_to_images(pptx_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert PowerPoint (PPTX) to images
    First converts to PDF, then PDF to images
    Returns tuple: (list of image paths, is_multipage)
    """
    try:
        # First convert PPTX to PDF
        pdf_path = await convert_pptx_to_pdf(pptx_path)

        try:
            # Then convert PDF to images
            return await convert_pdf_to_images(pdf_path, format, dpi)
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
    except Exception as e:
        raise ValueError(f"PPTX to {format.upper()} conversion failed: {str(e)}")


async def convert_ppt_to_images(ppt_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert PowerPoint (PPT) to images
    First converts to PDF, then PDF to images
    Returns tuple: (list of image paths, is_multipage)
    """
    try:
        # First convert PPT to PDF
        pdf_path = await convert_ppt_to_pdf(ppt_path)

        try:
            # Then convert PDF to images
            return await convert_pdf_to_images(pdf_path, format, dpi)
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
    except Exception as e:
        raise ValueError(f"PPT to {format.upper()} conversion failed: {str(e)}")


async def convert_xlsx_to_images(xlsx_path: str, format: str = 'png', dpi: int = 150) -> tuple:
    """
    Convert Excel (XLSX) to images
    First converts to PDF, then PDF to images
    Returns tuple: (list of image paths, is_multipage)
    """
    try:
        # First convert XLSX to PDF
        pdf_path = await convert_xlsx_to_pdf(xlsx_path)

        try:
            # Then convert PDF to images
            return await convert_pdf_to_images(pdf_path, format, dpi)
        finally:
            # Clean up temporary PDF
            if os.path.exists(pdf_path):
                os.unlink(pdf_path)
    except Exception as e:
        raise ValueError(f"XLSX to {format.upper()} conversion failed: {str(e)}")


# ==================== Check LibreOffice Status ====================

async def check_libreoffice_status() -> dict:
    """Check LibreOffice installation status"""
    result = {
        "available": False,
        "version": None,
        "supported_formats": []
    }

    try:
        proc = await asyncio.create_subprocess_exec(
            'libreoffice', '--version',
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        stdout, stderr = await proc.communicate()

        if proc.returncode == 0:
            result["available"] = True
            result["version"] = stdout.decode().strip().split('\n')[0]
            result["supported_formats"] = ['xlsx', 'xls', 'pptx', 'ppt', 'docx', 'doc', 'odt', 'ods', 'odp']
    except FileNotFoundError:
        pass

    return result
