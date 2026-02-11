from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageTemplate, Frame
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from io import BytesIO
from datetime import datetime
from typing import List, Dict, Any
import os
import logging

# DOCX generation imports
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)
import logging

logger = logging.getLogger(__name__)




def add_watermark_to_page(canvas_obj, doc):
    """Add watermark to page using onPage callback"""
    print("===== WATERMARK FUNCTION CALLED =====", flush=True)
    canvas_obj.saveState()
    canvas_obj.setFont("Helvetica-Bold", 60)
    canvas_obj.setFillAlpha(0.1)  # Reduced opacity for subtle watermark
    canvas_obj.setFillColor(colors.lightgrey)  # Light grey color
    
    width, height = letter
    canvas_obj.translate(width / 2, height / 2)
    canvas_obj.rotate(45)
    canvas_obj.drawCentredString(0, 0, "CalNConvert - FREE TIER")
    print("Watermark drawn on page", flush=True)
    
    canvas_obj.restoreState()


class WatermarkCanvas(canvas.Canvas):
    """Custom Canvas class that adds watermark to pages"""

    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.watermark_text = "CalNConvert - FREE TIER"

    def showPage(self):
        """Add watermark before showing page"""
        logger.info("===== WATERMARK: showPage() called =====")
        logger.info(f"Watermark text: {self.watermark_text}")
        
        # Save the current state
        self.saveState()

        # Set watermark properties
        self.setFont("Helvetica-Bold", 60)
        self.setFillAlpha(0.3)  # Increased to 0.3 for better visibility
        self.setFillColor(colors.HexColor('#EF4444'))  # Red color matching CalNConvert theme

        # Rotate and position watermark
        width, height = letter
        logger.info(f"Page dimensions: {width} x {height}")
        self.translate(width / 2, height / 2)
        self.rotate(45)
        self.drawCentredString(0, 0, self.watermark_text)
        logger.info("Watermark drawn successfully")

        # Restore state
        self.restoreState()

        # Call parent to show page
        super().showPage()
        logger.info("===== WATERMARK: showPage() completed =====")


async def generate_invoice_pdf(invoice_data: Dict, tier: str = "FREE") -> bytes:
    print(f">>> GENERATING INVOICE PDF (Tier: {tier}) <<<", flush=True)
    """
    Generate a PDF invoice with optional watermark based on tier.

    Args:
        invoice_data: Dictionary containing invoice details
        tier: User's subscription tier (FREE, STARTER, PRO)

    Returns:
        PDF content as bytes
    """
    buffer = BytesIO()

    # Only add watermark for FREE tier
    use_watermark = tier.upper() == "FREE"

    # Create PDF document with conditional watermark
    doc_kwargs = {
        "pagesize": letter,
        "rightMargin": 0.5 * inch,
        "leftMargin": 0.5 * inch,
        "topMargin": 0.5 * inch,
        "bottomMargin": 0.5 * inch,
    }

    if use_watermark:
        doc_kwargs["canvasmaker"] = WatermarkCanvas

    doc = SimpleDocTemplate(buffer, **doc_kwargs)

    elements = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#EF4444'),
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    elements.append(Paragraph("INVOICE", title_style))

    # Invoice header info
    invoice_info = f"""
    <b>Invoice #:</b> {invoice_data.get('invoice_number', 'N/A')}<br/>
    <b>Date:</b> {invoice_data.get('invoice_date', 'N/A')}<br/>
    <b>Due Date:</b> {invoice_data.get('due_date', 'N/A')}
    """
    elements.append(Paragraph(invoice_info, styles['Normal']))
    elements.append(Spacer(1, 0.2 * inch))

    # From/To section - Use Paragraph elements for proper HTML rendering
    cell_style = ParagraphStyle(
        'CellStyle',
        parent=styles['Normal'],
        fontSize=9,
    )

    from_info = Paragraph(
        f"""<b>{invoice_data.get('from_company', '')}</b><br/>
        {invoice_data.get('from_address', '')}<br/>
        Email: {invoice_data.get('from_email', '')}<br/>
        Phone: {invoice_data.get('from_phone', '')}""",
        cell_style
    )

    to_info = Paragraph(
        f"""<b>{invoice_data.get('to_client', '')}</b><br/>
        {invoice_data.get('to_address', '')}<br/>
        Email: {invoice_data.get('to_email', '')}""",
        cell_style
    )

    from_to_data = [
        ["FROM", "TO"],
        [from_info, to_info]
    ]

    # Standardized width: 7.5 inches (page width 8.5" - 0.5" margins on each side)
    from_to_table = Table(from_to_data, colWidths=[3.75 * inch, 3.75 * inch])
    from_to_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FEE2E2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))

    elements.append(from_to_table)
    elements.append(Spacer(1, 0.3 * inch))

    # Line items table - Standardized to 7.5 inches width
    line_items_data = [["Description", "Qty", "Rate", "Amount"]]

    for item in invoice_data.get('line_items', []):
        amount = item.get('quantity', 0) * item.get('rate', 0)
        line_items_data.append([
            item.get('description', ''),
            str(item.get('quantity', 0)),
            f"${item.get('rate', 0):.2f}",
            f"${amount:.2f}"
        ])

    line_items_table = Table(line_items_data, colWidths=[4 * inch, 1.5 * inch, 1 * inch, 1 * inch])
    line_items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FEE2E2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
    ]))

    elements.append(line_items_table)
    elements.append(Spacer(1, 0.2 * inch))

    # Totals section - Properly formatted table with consistent width and grid borders
    subtotal = sum(item.get('quantity', 0) * item.get('rate', 0) for item in invoice_data.get('line_items', []))
    discount_amount = subtotal * (invoice_data.get('discount', 0) / 100)
    taxable_amount = subtotal - discount_amount
    tax_amount = taxable_amount * (invoice_data.get('tax_rate', 0) / 100)
    total = taxable_amount + tax_amount

    totals_data = [
        ["SUBTOTAL", f"${subtotal:.2f}"],
    ]

    if invoice_data.get('discount', 0) > 0:
        totals_data.append([f"DISCOUNT ({invoice_data.get('discount', 0)}%)", f"-${discount_amount:.2f}"])

    if invoice_data.get('tax_rate', 0) > 0:
        totals_data.append([f"TAX ({invoice_data.get('tax_rate', 0)}%)", f"${tax_amount:.2f}"])

    totals_data.append(["TOTAL", f"${total:.2f}"])

    # Totals table - Same 7.5 inch width with proper grid borders (2 columns)
    totals_table = Table(totals_data, colWidths=[6.5 * inch, 1 * inch])
    totals_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -2), 10),
        ('FONTSIZE', (0, -1), (-1, -1), 12),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        # Grid borders for all rows
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        # Special styling for TOTAL row
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#FEE2E2')),
        ('TEXTCOLOR', (0, -1), (-1, -1), colors.HexColor('#991B1B')),
        ('BOTTOMPADDING', (0, -1), (-1, -1), 12),
    ]))

    elements.append(totals_table)
    elements.append(Spacer(1, 0.3 * inch))

    # Notes section
    if invoice_data.get('notes'):
        notes_style = ParagraphStyle(
            'Notes',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
        )
        elements.append(Paragraph(f"<b>Notes:</b> {invoice_data.get('notes')}", notes_style))

    # Build PDF with conditional watermark callbacks
    if use_watermark:
        doc.build(elements, onFirstPage=add_watermark_to_page, onLaterPages=add_watermark_to_page)
        print("##### Invoice PDF built with watermark (FREE tier) #####", flush=True)
    else:
        doc.build(elements)
        print(f"##### Invoice PDF built without watermark ({tier} tier) #####", flush=True)

    buffer.seek(0)
    return buffer.getvalue()


async def generate_nda_pdf(nda_data: Dict) -> bytes:
    """
    Generate a Non-Disclosure Agreement PDF.

    Args:
        nda_data: Dictionary containing NDA details

    Returns:
        PDF content as bytes
    """
    buffer = BytesIO()

    # Create PDF document (no watermark for Pro tier document)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    elements = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#1F2937'),
        spaceAfter=20,
        alignment=1,  # Center
        fontName='Helvetica-Bold'
    )
    elements.append(Paragraph("NON-DISCLOSURE AGREEMENT", title_style))
    elements.append(Spacer(1, 0.3 * inch))

    # Intro paragraph style
    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=12,
        alignment=4,  # Justify
    )

    heading_style = ParagraphStyle(
        'HeadingStyle',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=colors.HexColor('#374151'),
        spaceBefore=16,
        spaceAfter=8,
        fontName='Helvetica-Bold'
    )

    # Effective date
    effective_date = nda_data.get('effective_date', datetime.now().strftime('%B %d, %Y'))
    intro = f"""
    This Non-Disclosure Agreement (the "Agreement") is entered into as of <b>{effective_date}</b>
    (the "Effective Date") by and between:
    """
    elements.append(Paragraph(intro, body_style))
    elements.append(Spacer(1, 0.1 * inch))

    # Party information
    disclosing_party = nda_data.get('disclosing_party_name', 'Disclosing Party')
    disclosing_address = nda_data.get('disclosing_party_address', '')
    receiving_party = nda_data.get('receiving_party_name', 'Receiving Party')
    receiving_address = nda_data.get('receiving_party_address', '')

    party_info = f"""
    <b>Disclosing Party:</b> {disclosing_party}<br/>
    Address: {disclosing_address}<br/><br/>
    <b>Receiving Party:</b> {receiving_party}<br/>
    Address: {receiving_address}
    """
    elements.append(Paragraph(party_info, body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Purpose
    purpose = nda_data.get('purpose', 'business discussions and potential collaboration')
    elements.append(Paragraph("1. PURPOSE", heading_style))
    purpose_text = f"""
    The Disclosing Party wishes to disclose certain confidential and proprietary information
    to the Receiving Party for the purpose of: <b>{purpose}</b>.
    """
    elements.append(Paragraph(purpose_text, body_style))

    # Definition of Confidential Information
    elements.append(Paragraph("2. DEFINITION OF CONFIDENTIAL INFORMATION", heading_style))
    confidential_def = """
    "Confidential Information" means any and all non-public information, including but not limited to:
    technical data, trade secrets, know-how, research, product plans, products, services, customers,
    customer lists, markets, software, developments, inventions, processes, formulas, technology,
    designs, drawings, engineering, hardware configuration information, marketing, finances, or
    other business information disclosed by the Disclosing Party to the Receiving Party, either
    directly or indirectly, in writing, orally, or by inspection of tangible objects.
    """
    elements.append(Paragraph(confidential_def, body_style))

    # Obligations
    elements.append(Paragraph("3. OBLIGATIONS OF RECEIVING PARTY", heading_style))
    obligations = """
    The Receiving Party agrees to:<br/>
    a) Hold and maintain the Confidential Information in strict confidence;<br/>
    b) Not disclose the Confidential Information to any third parties without prior written consent;<br/>
    c) Not use the Confidential Information for any purpose other than as stated in this Agreement;<br/>
    d) Protect the Confidential Information using the same degree of care used to protect its own
    confidential information, but in no event less than reasonable care.
    """
    elements.append(Paragraph(obligations, body_style))

    # Term
    duration_years = nda_data.get('duration_years', 2)
    elements.append(Paragraph("4. TERM", heading_style))
    term_text = f"""
    This Agreement shall remain in effect for a period of <b>{duration_years} year(s)</b> from the
    Effective Date, unless terminated earlier by either party with thirty (30) days written notice.
    The obligations of confidentiality shall survive termination of this Agreement.
    """
    elements.append(Paragraph(term_text, body_style))

    # Return of Information
    elements.append(Paragraph("5. RETURN OF INFORMATION", heading_style))
    return_text = """
    Upon termination of this Agreement or upon request by the Disclosing Party, the Receiving Party
    shall promptly return or destroy all Confidential Information and any copies thereof.
    """
    elements.append(Paragraph(return_text, body_style))

    # Governing Law
    governing_law = nda_data.get('governing_law', 'the State of [Your State]')
    elements.append(Paragraph("6. GOVERNING LAW", heading_style))
    law_text = f"""
    This Agreement shall be governed by and construed in accordance with the laws of <b>{governing_law}</b>,
    without regard to its conflict of law provisions.
    """
    elements.append(Paragraph(law_text, body_style))

    # Miscellaneous
    elements.append(Paragraph("7. MISCELLANEOUS", heading_style))
    misc_text = """
    This Agreement constitutes the entire agreement between the parties with respect to the subject
    matter hereof and supersedes all prior negotiations, representations, or agreements relating thereto.
    This Agreement may not be amended except by a written instrument signed by both parties.
    """
    elements.append(Paragraph(misc_text, body_style))
    elements.append(Spacer(1, 0.4 * inch))

    # Signature section
    elements.append(Paragraph("<b>IN WITNESS WHEREOF</b>, the parties have executed this Agreement as of the date first written above.", body_style))
    elements.append(Spacer(1, 0.4 * inch))

    # Signature table
    sig_style = ParagraphStyle(
        'SigStyle',
        parent=styles['Normal'],
        fontSize=10,
    )

    sig_data = [
        [Paragraph("<b>DISCLOSING PARTY:</b>", sig_style), Paragraph("<b>RECEIVING PARTY:</b>", sig_style)],
        ["", ""],
        ["_________________________", "_________________________"],
        [f"Name: {disclosing_party}", f"Name: {receiving_party}"],
        ["Date: _________________", "Date: _________________"],
    ]

    sig_table = Table(sig_data, colWidths=[3.5 * inch, 3.5 * inch])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))

    elements.append(sig_table)

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


async def generate_employment_letter_pdf(letter_data: Dict) -> bytes:
    """
    Generate an Employment Offer/Confirmation Letter PDF.

    Args:
        letter_data: Dictionary containing letter details

    Returns:
        PDF content as bytes
    """
    buffer = BytesIO()

    # Create PDF document (no watermark for Pro tier document)
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    elements = []
    styles = getSampleStyleSheet()

    # Company header style
    header_style = ParagraphStyle(
        'HeaderStyle',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#374151'),
        alignment=0,  # Left
    )

    body_style = ParagraphStyle(
        'BodyText',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=12,
    )

    # Company letterhead
    company_name = letter_data.get('company_name', 'Company Name')
    company_address = letter_data.get('company_address', '')
    company_phone = letter_data.get('company_phone', '')
    company_email = letter_data.get('company_email', '')

    letterhead = f"""
    <b>{company_name}</b><br/>
    {company_address}<br/>
    Phone: {company_phone}<br/>
    Email: {company_email}
    """
    elements.append(Paragraph(letterhead, header_style))
    elements.append(Spacer(1, 0.3 * inch))

    # Date
    letter_date = letter_data.get('letter_date', datetime.now().strftime('%B %d, %Y'))
    elements.append(Paragraph(letter_date, body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Recipient
    employee_name = letter_data.get('employee_name', 'Employee Name')
    employee_address = letter_data.get('employee_address', '')

    recipient = f"""
    {employee_name}<br/>
    {employee_address}
    """
    elements.append(Paragraph(recipient, body_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Salutation
    elements.append(Paragraph(f"Dear {employee_name},", body_style))
    elements.append(Spacer(1, 0.1 * inch))

    # Letter type
    letter_type = letter_data.get('letter_type', 'offer')  # 'offer' or 'confirmation'

    if letter_type == 'offer':
        # Offer Letter Content
        position = letter_data.get('position', 'Position Title')
        department = letter_data.get('department', 'Department')
        start_date = letter_data.get('start_date', 'TBD')
        salary = letter_data.get('salary', 0)
        salary_frequency = letter_data.get('salary_frequency', 'annually')
        manager = letter_data.get('reporting_manager', 'Manager Name')

        intro = f"""
        We are pleased to offer you the position of <b>{position}</b> in the <b>{department}</b> department
        at {company_name}. We believe your skills and experience will be a valuable addition to our team.
        """
        elements.append(Paragraph(intro, body_style))

        details = f"""
        <b>Position Details:</b><br/>
        • Position: {position}<br/>
        • Department: {department}<br/>
        • Start Date: {start_date}<br/>
        • Reporting To: {manager}<br/>
        • Compensation: ${salary:,.2f} {salary_frequency}
        """
        elements.append(Paragraph(details, body_style))

        # Benefits if provided
        benefits = letter_data.get('benefits', '')
        if benefits:
            benefits_text = f"""
            <b>Benefits:</b><br/>
            {benefits}
            """
            elements.append(Paragraph(benefits_text, body_style))

        conditions = """
        This offer is contingent upon satisfactory completion of background verification and
        reference checks. Please confirm your acceptance of this offer by signing and returning
        this letter prior to your start date listed above.
        """
        elements.append(Paragraph(conditions, body_style))

    else:
        # Employment Confirmation Letter
        position = letter_data.get('position', 'Position Title')
        department = letter_data.get('department', 'Department')
        hire_date = letter_data.get('hire_date', 'Date')
        salary = letter_data.get('salary', 0)
        salary_frequency = letter_data.get('salary_frequency', 'annually')

        intro = f"""
        This letter confirms your employment with {company_name} in the position of <b>{position}</b>
        in the <b>{department}</b> department.
        """
        elements.append(Paragraph(intro, body_style))

        details = f"""
        <b>Employment Details:</b><br/>
        • Position: {position}<br/>
        • Department: {department}<br/>
        • Date of Hire: {hire_date}<br/>
        • Current Salary: ${salary:,.2f} {salary_frequency}<br/>
        • Employment Status: {letter_data.get('employment_status', 'Full-time')}
        """
        elements.append(Paragraph(details, body_style))

        confirmation = """
        If you require any additional information regarding your employment, please do not
        hesitate to contact the Human Resources department.
        """
        elements.append(Paragraph(confirmation, body_style))

    elements.append(Spacer(1, 0.3 * inch))

    # Additional terms if provided
    additional_terms = letter_data.get('additional_terms', '')
    if additional_terms:
        elements.append(Paragraph(f"<b>Additional Terms:</b><br/>{additional_terms}", body_style))
        elements.append(Spacer(1, 0.2 * inch))

    # Closing
    elements.append(Paragraph("Sincerely,", body_style))
    elements.append(Spacer(1, 0.4 * inch))

    hr_name = letter_data.get('hr_name', 'HR Manager')
    hr_title = letter_data.get('hr_title', 'Human Resources')

    closing = f"""
    _________________________<br/>
    {hr_name}<br/>
    {hr_title}<br/>
    {company_name}
    """
    elements.append(Paragraph(closing, body_style))

    if letter_type == 'offer':
        elements.append(Spacer(1, 0.4 * inch))
        acceptance = f"""
        <b>ACCEPTANCE</b><br/><br/>
        I, {employee_name}, accept the offer of employment as described above.<br/><br/>
        Signature: _________________________<br/>
        Date: _________________________
        """
        elements.append(Paragraph(acceptance, body_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


async def generate_receipt_pdf(receipt_data: Dict, tier: str = "FREE") -> bytes:
    print(f">>> GENERATING RECEIPT PDF (Tier: {tier}) <<<", flush=True)
    """
    Generate a PDF receipt with optional watermark based on tier.

    Args:
        receipt_data: Dictionary containing receipt details
        tier: User's subscription tier (FREE, STARTER, PRO)

    Returns:
        PDF content as bytes
    """
    buffer = BytesIO()

    # Only add watermark for FREE tier
    use_watermark = tier.upper() == "FREE"

    # Create PDF document with conditional watermark
    doc_kwargs = {
        "pagesize": letter,
        "rightMargin": 0.5 * inch,
        "leftMargin": 0.5 * inch,
        "topMargin": 0.5 * inch,
        "bottomMargin": 0.5 * inch,
    }

    if use_watermark:
        doc_kwargs["canvasmaker"] = WatermarkCanvas

    doc = SimpleDocTemplate(buffer, **doc_kwargs)

    elements = []
    styles = getSampleStyleSheet()

    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#EF4444'),
        spaceAfter=20,
        alignment=1  # Center alignment
    )
    elements.append(Paragraph("RECEIPT", title_style))

    # Receipt header info
    receipt_info = f"""
    <b>Receipt #:</b> {receipt_data.get('receipt_number', 'N/A')}<br/>
    <b>Date:</b> {receipt_data.get('receipt_date', 'N/A')}<br/>
    <b>Payment Method:</b> {receipt_data.get('payment_method', 'Cash')}
    """
    elements.append(Paragraph(receipt_info, styles['Normal']))
    elements.append(Spacer(1, 0.2 * inch))

    # Business info
    business_info = f"""
    <b>{receipt_data.get('from_business', '')}</b><br/>
    {receipt_data.get('from_address', '')}<br/>
    Phone: {receipt_data.get('from_phone', '')}<br/>
    Email: {receipt_data.get('from_email', '')}
    """
    business_style = ParagraphStyle(
        'BusinessInfo',
        parent=styles['Normal'],
        fontSize=9,
        textColor=colors.grey,
    )
    elements.append(Paragraph(business_info, business_style))
    elements.append(Spacer(1, 0.2 * inch))

    # Customer name if provided
    if receipt_data.get('customer_name'):
        elements.append(Paragraph(f"<b>Customer:</b> {receipt_data.get('customer_name')}", styles['Normal']))
        elements.append(Spacer(1, 0.1 * inch))

    # Line items table - Standardized to 7.5 inches width
    line_items_data = [["Description", "Qty", "Rate", "Amount"]]

    for item in receipt_data.get('line_items', []):
        amount = item.get('quantity', 0) * item.get('rate', 0)
        line_items_data.append([
            item.get('description', ''),
            str(item.get('quantity', 0)),
            f"${item.get('rate', 0):.2f}",
            f"${amount:.2f}"
        ])

    line_items_table = Table(line_items_data, colWidths=[4 * inch, 1.5 * inch, 1 * inch, 1 * inch])
    line_items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FEE2E2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#FAFAFA')]),
    ]))

    elements.append(line_items_table)
    elements.append(Spacer(1, 0.2 * inch))

    # Total - Standardized to 7.5 inches width with grid borders
    total = sum(item.get('quantity', 0) * item.get('rate', 0) for item in receipt_data.get('line_items', []))

    totals_data = [
        ["TOTAL", f"${total:.2f}"]
    ]

    totals_table = Table(totals_data, colWidths=[6.5 * inch, 1 * inch])
    totals_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (1, 0), 12),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FEE2E2')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#991B1B')),
        ('GRID', (0, 0), (-1, 0), 1, colors.grey),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
    ]))

    elements.append(totals_table)

    # Build PDF with conditional watermark callbacks
    if use_watermark:
        doc.build(elements, onFirstPage=add_watermark_to_page, onLaterPages=add_watermark_to_page)
        print("##### Receipt PDF built with watermark (FREE tier) #####", flush=True)
    else:
        doc.build(elements)
        print(f"##### Receipt PDF built without watermark ({tier} tier) #####", flush=True)

    buffer.seek(0)
    return buffer.getvalue()


# ==================== DOCX Export Functions (Pro Tier) ====================

async def generate_invoice_docx(invoice_data: Dict) -> bytes:
    """
    Generate an Invoice in DOCX format (Pro tier only).

    Args:
        invoice_data: Dictionary containing invoice details

    Returns:
        DOCX content as bytes
    """
    if not DOCX_AVAILABLE:
        raise ValueError("DOCX generation not available. Please install python-docx.")

    doc = Document()

    # Title
    title = doc.add_heading('INVOICE', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Invoice info
    doc.add_paragraph(f"Invoice #: {invoice_data.get('invoice_number', 'N/A')}")
    doc.add_paragraph(f"Date: {invoice_data.get('invoice_date', 'N/A')}")
    doc.add_paragraph(f"Due Date: {invoice_data.get('due_date', 'N/A')}")
    doc.add_paragraph()

    # From section
    doc.add_heading('From:', level=2)
    doc.add_paragraph(f"{invoice_data.get('from_company', '')}")
    doc.add_paragraph(f"{invoice_data.get('from_address', '')}")
    doc.add_paragraph(f"Email: {invoice_data.get('from_email', '')}")
    doc.add_paragraph(f"Phone: {invoice_data.get('from_phone', '')}")

    # To section
    doc.add_heading('To:', level=2)
    doc.add_paragraph(f"{invoice_data.get('to_client', '')}")
    doc.add_paragraph(f"{invoice_data.get('to_address', '')}")
    doc.add_paragraph(f"Email: {invoice_data.get('to_email', '')}")
    doc.add_paragraph()

    # Line items table
    doc.add_heading('Items:', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Description'
    header_cells[1].text = 'Qty'
    header_cells[2].text = 'Rate'
    header_cells[3].text = 'Amount'

    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add line items
    for item in invoice_data.get('line_items', []):
        amount = item.get('quantity', 0) * item.get('rate', 0)
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('description', '')
        row_cells[1].text = str(item.get('quantity', 0))
        row_cells[2].text = f"${item.get('rate', 0):.2f}"
        row_cells[3].text = f"${amount:.2f}"

    doc.add_paragraph()

    # Totals
    subtotal = sum(item.get('quantity', 0) * item.get('rate', 0) for item in invoice_data.get('line_items', []))
    discount_amount = subtotal * (invoice_data.get('discount', 0) / 100)
    taxable_amount = subtotal - discount_amount
    tax_amount = taxable_amount * (invoice_data.get('tax_rate', 0) / 100)
    total = taxable_amount + tax_amount

    doc.add_paragraph(f"Subtotal: ${subtotal:.2f}")
    if invoice_data.get('discount', 0) > 0:
        doc.add_paragraph(f"Discount ({invoice_data.get('discount', 0)}%): -${discount_amount:.2f}")
    if invoice_data.get('tax_rate', 0) > 0:
        doc.add_paragraph(f"Tax ({invoice_data.get('tax_rate', 0)}%): ${tax_amount:.2f}")

    total_para = doc.add_paragraph()
    total_run = total_para.add_run(f"TOTAL: ${total:.2f}")
    total_run.bold = True

    # Notes
    if invoice_data.get('notes'):
        doc.add_paragraph()
        doc.add_paragraph(f"Notes: {invoice_data.get('notes')}")

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


async def generate_receipt_docx(receipt_data: Dict) -> bytes:
    """
    Generate a Receipt in DOCX format (Pro tier only).

    Args:
        receipt_data: Dictionary containing receipt details

    Returns:
        DOCX content as bytes
    """
    if not DOCX_AVAILABLE:
        raise ValueError("DOCX generation not available. Please install python-docx.")

    doc = Document()

    # Title
    title = doc.add_heading('RECEIPT', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Receipt info
    doc.add_paragraph(f"Receipt #: {receipt_data.get('receipt_number', 'N/A')}")
    doc.add_paragraph(f"Date: {receipt_data.get('receipt_date', 'N/A')}")
    doc.add_paragraph(f"Payment Method: {receipt_data.get('payment_method', 'Cash')}")
    doc.add_paragraph()

    # Business info
    doc.add_heading('From:', level=2)
    doc.add_paragraph(f"{receipt_data.get('from_business', '')}")
    doc.add_paragraph(f"{receipt_data.get('from_address', '')}")
    doc.add_paragraph(f"Phone: {receipt_data.get('from_phone', '')}")
    doc.add_paragraph(f"Email: {receipt_data.get('from_email', '')}")

    # Customer info
    if receipt_data.get('customer_name'):
        doc.add_paragraph(f"Customer: {receipt_data.get('customer_name')}")
    doc.add_paragraph()

    # Line items table
    doc.add_heading('Items:', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Description'
    header_cells[1].text = 'Qty'
    header_cells[2].text = 'Rate'
    header_cells[3].text = 'Amount'

    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add line items
    for item in receipt_data.get('line_items', []):
        amount = item.get('quantity', 0) * item.get('rate', 0)
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('description', '')
        row_cells[1].text = str(item.get('quantity', 0))
        row_cells[2].text = f"${item.get('rate', 0):.2f}"
        row_cells[3].text = f"${amount:.2f}"

    doc.add_paragraph()

    # Total
    total = sum(item.get('quantity', 0) * item.get('rate', 0) for item in receipt_data.get('line_items', []))
    total_para = doc.add_paragraph()
    total_run = total_para.add_run(f"TOTAL: ${total:.2f}")
    total_run.bold = True

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
